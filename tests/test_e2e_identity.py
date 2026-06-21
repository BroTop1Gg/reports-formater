#!/usr/bin/env python3
"""
E2E Structural Identity Test.

Proves that the chunked MCP-driven pipeline produces identical output
to the legacy monolithic YAML engine.
"""

import unittest
import tempfile
import shutil
from pathlib import Path

import yaml
from docx import Document

from src.report_factory import ReportFactory
from src.sdk.session import ReportSession


class TestStructuralIdentity(unittest.TestCase):
    """
    Verify that monolithic compilation and sequential chunk compilation
    produce structurally identical .docx output.
    """

    def setUp(self):
        """Set up test fixtures."""
        self.project_root = Path(__file__).parent.parent
        self.test_yaml = self.project_root / "tests" / "input" / "test_with_title.yaml"
        self.temp_dir = tempfile.mkdtemp()
        self.temp_path = Path(self.temp_dir)

    def tearDown(self):
        """Clean up temporary files."""
        if self.temp_path.exists():
            shutil.rmtree(self.temp_path)

    def test_monolithic_vs_chunked_parity(self):
        """
        Compare monolithic ReportFactory.build() vs sequential ReportSession chunks.
        
        Workflow:
        1. Load reference YAML
        2. Compile monolithically -> monolithic_ref.docx
        3. Parse YAML into chunks
        4. Feed chunks to ReportSession -> chunked_ref.docx
        5. Assert complete structural parity
        """
        # Load YAML data
        with open(self.test_yaml, "r", encoding="utf-8") as f:
            yaml_data = yaml.safe_load(f)

        # === MONOLITHIC COMPILATION ===
        monolithic_output = self.temp_path / "monolithic_ref.docx"
        factory = ReportFactory()
        factory.build(
            yaml_data=yaml_data,
            output_path=monolithic_output,
            resource_path=self.test_yaml.parent
        )

        # === CHUNKED COMPILATION ===
        # Extract metadata and content
        metadata = yaml_data.get("metadata", {})
        content_nodes = yaml_data.get("content", [])

        # Initialize session with metadata
        session = ReportSession(metadata=metadata)

        # Feed content nodes as individual chunks (one node per chunk for granular testing)
        for node in content_nodes:
            # Convert node to YAML string
            chunk_yaml = yaml.dump([node], allow_unicode=True, default_flow_style=False)
            result = session.add_chunk(chunk_yaml)
            
            # Each chunk should succeed
            self.assertEqual(
                result["status"], "success",
                f"Chunk failed: {node.get('type', 'unknown')} - {result.get('errors', [])}"
            )

        # Finalize chunked output
        chunked_output = self.temp_path / "chunked_ref.docx"
        session.finalize(
            output_path=chunked_output,
            resource_path=self.test_yaml.parent
        )

        # === STRUCTURAL PARITY ASSERTIONS ===
        self._assert_document_parity(monolithic_output, chunked_output)

    def _assert_document_parity(self, doc1_path: Path, doc2_path: Path):
        """
        Assert complete structural parity between two .docx files.
        
        Checks:
        - Paragraph count and text content
        - Table count, row count, and cell values
        - Inline styles (bold, italic)
        - Spacing configurations
        """
        doc1 = Document(str(doc1_path))
        doc2 = Document(str(doc2_path))

        # 1. Paragraph count
        self.assertEqual(
            len(doc1.paragraphs), len(doc2.paragraphs),
            f"Paragraph count mismatch: {len(doc1.paragraphs)} vs {len(doc2.paragraphs)}"
        )

        # 2. Paragraph text content
        for i, (p1, p2) in enumerate(zip(doc1.paragraphs, doc2.paragraphs)):
            self.assertEqual(
                p1.text, p2.text,
                f"Paragraph {i} text mismatch:\n"
                f"  Monolithic: {p1.text[:100]}...\n"
                f"  Chunked:    {p2.text[:100]}..."
            )

        # 3. Table count
        self.assertEqual(
            len(doc1.tables), len(doc2.tables),
            f"Table count mismatch: {len(doc1.tables)} vs {len(doc2.tables)}"
        )

        # 4. Table structure and content
        for t_idx, (t1, t2) in enumerate(zip(doc1.tables, doc2.tables)):
            self.assertEqual(
                len(t1.rows), len(t2.rows),
                f"Table {t_idx} row count mismatch: {len(t1.rows)} vs {len(t2.rows)}"
            )

            for r_idx, (r1, r2) in enumerate(zip(t1.rows, t2.rows)):
                self.assertEqual(
                    len(r1.cells), len(r2.cells),
                    f"Table {t_idx} Row {r_idx} cell count mismatch"
                )

                for c_idx, (c1, c2) in enumerate(zip(r1.cells, r2.cells)):
                    self.assertEqual(
                        c1.text, c2.text,
                        f"Table {t_idx} Row {r_idx} Cell {c_idx} text mismatch:\n"
                        f"  Monolithic: {c1.text[:50]}\n"
                        f"  Chunked:    {c2.text[:50]}"
                    )

        # 5. Inline styles (bold, italic) - sample check
        # Check first 10 paragraphs with runs
        checked = 0
        for p1, p2 in zip(doc1.paragraphs, doc2.paragraphs):
            if len(p1.runs) > 0 and len(p2.runs) > 0:
                # Compare run-level formatting
                for r1, r2 in zip(p1.runs, p2.runs):
                    self.assertEqual(
                        r1.bold, r2.bold,
                        f"Bold mismatch in paragraph: {p1.text[:50]}"
                    )
                    self.assertEqual(
                        r1.italic, r2.italic,
                        f"Italic mismatch in paragraph: {p1.text[:50]}"
                    )
                
                checked += 1
                if checked >= 10:
                    break

        # 6. Paragraph alignment (sample check)
        for i, (p1, p2) in enumerate(zip(doc1.paragraphs[:20], doc2.paragraphs[:20])):
            self.assertEqual(
                p1.alignment, p2.alignment,
                f"Paragraph {i} alignment mismatch: {p1.alignment} vs {p2.alignment}"
            )

        # 7. Page layout (margins)
        for s1, s2 in zip(doc1.sections, doc2.sections):
            self.assertEqual(s1.top_margin, s2.top_margin, "Top margin mismatch")
            self.assertEqual(s1.bottom_margin, s2.bottom_margin, "Bottom margin mismatch")
            self.assertEqual(s1.left_margin, s2.left_margin, "Left margin mismatch")
            self.assertEqual(s1.right_margin, s2.right_margin, "Right margin mismatch")

    def test_chunked_session_accumulates_correctly(self):
        """
        Verify that ReportSession correctly accumulates nodes across multiple chunks.
        """
        # Load YAML
        with open(self.test_yaml, "r", encoding="utf-8") as f:
            yaml_data = yaml.safe_load(f)

        content_nodes = yaml_data.get("content", [])
        metadata = yaml_data.get("metadata", {})

        session = ReportSession(metadata=metadata)

        # Add all nodes as individual chunks
        for node in content_nodes:
            chunk_yaml = yaml.dump([node], allow_unicode=True)
            result = session.add_chunk(chunk_yaml)
            self.assertEqual(result["status"], "success")

        # Verify all nodes accumulated
        self.assertEqual(
            len(session.nodes), len(content_nodes),
            f"Expected {len(content_nodes)} nodes, got {len(session.nodes)}"
        )

        # Verify chunk index incremented correctly
        self.assertEqual(session.chunk_index, len(content_nodes))

    def test_monolithic_yaml_vs_mcp_markdown_parity(self):
        """
        Compare monolithic YAML compilation vs Markdown transpilation pipeline.
        
        Workflow:
        1. Load reference YAML and compile monolithically -> monolithic_ref.docx
        2. Load equivalent Markdown and compile via add_markdown_chunk -> markdown_mcp_ref.docx
        3. Assert complete structural parity between both outputs
        
        This test validates the entire Markdown transpiler bridge end-to-end.
        """
        # Load YAML data
        with open(self.test_yaml, "r", encoding="utf-8") as f:
            yaml_data = yaml.safe_load(f)

        # === MONOLITHIC YAML COMPILATION ===
        monolithic_output = self.temp_path / "monolithic_ref.docx"
        factory = ReportFactory()
        factory.build(
            yaml_data=yaml_data,
            output_path=monolithic_output,
            resource_path=self.test_yaml.parent
        )

        # === MARKDOWN TRANSPILATION COMPILATION ===
        # Load equivalent Markdown
        markdown_file = self.project_root / "tests" / "input" / "test_with_title.md"
        with open(markdown_file, "r", encoding="utf-8") as f:
            markdown_content = f.read()

        # Initialize session with same metadata as YAML
        metadata = yaml_data.get("metadata", {})
        session = ReportSession(metadata=metadata)

        # Feed entire Markdown content as single chunk
        result = session.add_markdown_chunk(markdown_content)
        
        # Should succeed
        self.assertEqual(
            result["status"], "success",
            f"Markdown transpilation failed: {result.get('errors', [])}"
        )

        # Finalize markdown output
        markdown_output = self.temp_path / "markdown_mcp_ref.docx"
        session.finalize(
            output_path=markdown_output,
            resource_path=self.test_yaml.parent
        )

        # === STRUCTURAL PARITY ASSERTIONS ===
        self._assert_document_parity(monolithic_output, markdown_output)


if __name__ == "__main__":
    unittest.main()
