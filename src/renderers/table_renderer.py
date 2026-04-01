"""
Table Renderer for Reports-Formater.

Renders grid tables with header row repetition support (DSTU requirement).
"""

import logging
from typing import List, Union

from docx.oxml import OxmlElement
from docx.shared import Pt, Cm

from src.utils.docx_utils import (
    add_table_borders,
    optimize_table_width_and_alignment,
    get_alignment_enum
)
from src.renderers.base import BaseRenderer, RenderContext
from src.config.schemas import TableData
from src.utils.formatting import parse_inline_formatting

logger = logging.getLogger(__name__)


class TableRenderer(BaseRenderer):
    """
    Renderer for table content nodes.

    Handles:
    - Grid table creation
    - Header row repetition (w:tblHeader OXML)
    - Cell content with inline formatting
    - Optional caption above table (DSTU requirement)
    - Nested content via dispatch
    """

    @property
    def node_type(self) -> str:
        """Return the node type this renderer handles."""
        return "table"

    def render(self, context: RenderContext, data: TableData) -> None:
        """
        Render table content to document.

        [IMPORTANT] DSTU 3008-2015 requires "Continued table ..." text on each new page
        when a table breaks across pages. However, automatic insertion of such labels
        can destabilize layout or require complex nested table structures which may
        violate our "Dumb Builder" philosophy. Currently, this is omitted to ensure
        document structural integrity.


        Args:
            context: Render context with container and config.
            data: Validated TableData model.
        """
        # todo: fix that issue.

        rows_data = data.rows
        if not rows_data:
            logger.warning("TableRenderer: Empty table data, skipping")
            return

        num_rows = len(rows_data)
        num_cols = len(rows_data[0]) if rows_data else 0

        if num_cols == 0:
            logger.warning("TableRenderer: No columns in first row, skipping")
            return

        # Render caption before table if provided
        if data.caption:
            self._render_caption(context, data.caption)

        # Create table
        table = context.container.add_table(rows=num_rows, cols=num_cols)
        table.autofit = False

        # Determine explicit width to avoid sub-millimeter offsets in LibreOffice/Word
        # If we cannot determine page width from sections, we will rely on optimize_table_width_and_alignment (100% width)
        # without setting explicit column width.
        item_width = None
        try:
            section = context.doc.sections[0]
            available = section.page_width - section.left_margin - section.right_margin
            if available > 0:
                item_width = available
        except Exception:
            pass

        # Distribute width equally among columns if available, otherwise let Word manage it
        if num_cols > 0 and item_width is not None:
            col_width = int(item_width / num_cols)
            for col in table.columns:
                col.width = col_width

        # Apply table style
        resolved_style = context.style_manager.get_style_name(
            data.style,
            fallback="Table Grid"
        )
        try:
            table.style = resolved_style
        except (KeyError, ValueError) as e:
            logger.debug(f"Could not apply table style '{resolved_style}': {e}")

        # Optimize spacing and layout
        optimize_table_width_and_alignment(table)

        # Explicitly apply grid borders.
        # LibreOffice templates often lack the 'Table Grid' style and fall back to 'Normal',
        # causing borders to vanish without this explicit XML injection.
        add_table_borders(table)

        # Get font config
        fonts = context.config.fonts
        styles = context.config.styles
        normal_size = styles.normal.font_size_pt

        # Fill table cells
        for r_idx, row_items in enumerate(rows_data):
            # First row gets header marker for repeat on page break (if enabled)
            if r_idx == 0 and data.repeat_header:
                self._mark_header_row(table.rows[0])

            row_cells = table.rows[r_idx].cells

            for c_idx, cell_content in enumerate(row_items):
                if c_idx >= len(row_cells):
                    continue

                cell = row_cells[c_idx]
                cell_p = cell.paragraphs[0]

                # Prevent default paragraph indentation (like first-line indent) in table cells
                pf = cell_p.paragraph_format
                pf.first_line_indent = Pt(0)
                pf.left_indent = Pt(0)
                pf.right_indent = Pt(0)

                # Handle content (simple string or dict for nested)
                if isinstance(cell_content, dict):
                    # Dispatch for nested content rendering
                    cell_context = context.with_container(cell)
                    context.dispatch(cell_content)
                else:
                    # Multi-paragraph support for table cells (Requirement: \n -> Paragraph)
                    content_str = str(cell_content)
                    lines = content_str.split('\n')
                    
                    for i, line in enumerate(lines):
                        if i == 0:
                            # Re-use the existing first paragraph
                            p = cell.paragraphs[0]
                        else:
                            # Append additional paragraphs to the cell
                            p = cell.add_paragraph()
                            
                        # Ensure all paragraphs within the cell have consistent padding/indents
                        # Table cells usually have 0 indents in DSTU.
                        pf = p.paragraph_format
                        pf.first_line_indent = Pt(0)
                        pf.left_indent = Pt(0)
                        pf.right_indent = Pt(0)
                        
                        # Apply spacing from global 'normal' style or cell defaults
                        pf.line_spacing = styles.normal.line_spacing
                        pf.space_before = Pt(styles.normal.space_before_pt / 2) # Reduced for tables
                        pf.space_after = Pt(styles.normal.space_after_pt / 2)
                        
                        parse_inline_formatting(
                            paragraph=p,
                            text=line,
                            default_font=fonts.default_name,
                            custom_font=styles.normal.font_name, # Tables usually follow normal font
                            code_font=styles.inline_code.font_name if styles.inline_code else fonts.code_name,
                            base_size_pt=normal_size,
                        )

    def _render_caption(self, context: RenderContext, caption: str) -> None:
        """
        Render caption paragraph before the table.

        The caption formatting follows settings defined in 'caption' section of
        report_styles.json. This approach ensures total separation of concerns
        between Python logic and visual representation.

        Args:
            context: Current render context providing access to doc and config.
            caption: Raw caption text from YAML (supports inline formatting).
        """
        p = context.container.add_paragraph()

        # Access centralized style configuration
        # This eliminates "magic numbers" in code and respects the configuration hierarchy.
        style_config = context.config.styles.caption
        styles = context.config.styles
        fonts = context.config.fonts

        # Apply paragraph-level formatting from config
        pf = p.paragraph_format
        pf.alignment = get_alignment_enum(style_config.alignment)
        pf.line_spacing = style_config.line_spacing
        pf.first_line_indent = Cm(style_config.first_line_indent_cm)
        pf.space_before = Pt(style_config.space_before_pt)
        pf.space_after = Pt(style_config.space_after_pt)
        pf.keep_with_next = True

        # Reset indents to 0 unless specified otherwise to avoid inherited indents
        pf.left_indent = Pt(0)
        pf.right_indent = Pt(0)

        # Render text with support for inline markdown (**bold**, *italic*, `code`)
        parse_inline_formatting(
            paragraph=p,
            text=caption,
            default_font=fonts.default_name,
            custom_font=style_config.font_name,
            code_font=styles.inline_code.font_name if styles.inline_code else fonts.code_name,
            base_size_pt=style_config.font_size_pt
        )

    def _mark_header_row(self, row) -> None:
        """
        Mark row as header for repeat on page break.

        Uses OXML w:tblHeader element.
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)
