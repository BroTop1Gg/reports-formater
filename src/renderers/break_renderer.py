"""
Break Renderer for Reports-Formater.

Renders explicit line breaks (vertical spacing) and page breaks.
"""

import logging
from docx.enum.text import WD_BREAK

from src.renderers.base import BaseRenderer, RenderContext
from src.config.schemas import BreakData

logger = logging.getLogger(__name__)


class BreakRenderer(BaseRenderer):
    """
    Renderer for break nodes.
    
    Handles:
    - Line breaks (empty paragraphs for vertical spacing)
    - Page breaks (hard page breaks)
    """
    
    @property
    def node_type(self) -> str:
        """Return the node type this renderer handles."""
        return "break"
    
    def render(self, context: RenderContext, data: BreakData) -> None:
        """
        Render break to document.
        
        Args:
            context: Render context.
            data: Validated BreakData model.
        """
        if data.style == "page":
            # Add page break
            # We add a paragraph with a page break run
            p = context.container.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
            logger.debug("BreakRenderer: Inserted PAGE break")
            
        elif data.style == "line":
            # Add vertical spacing
            # We treat this as adding 'count' empty paragraphs
            # This is "explicit new line" as requested
            count = data.count
            for _ in range(count):
                context.container.add_paragraph()
            logger.debug(f"BreakRenderer: Inserted {count} LINE break(s)")
            
        elif data.style == "section":
            # Section breaks are complex in python-docx as they are tied to creating new sections
            # For now, treat as page break or implement if needed
            logger.warning("BreakRenderer: Section breaks not fully implemented, using page break.")
            p = context.container.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
