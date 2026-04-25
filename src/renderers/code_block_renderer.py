"""
Code Block Renderer for Reports-Formater.

Renders formatted code blocks with optional captions.
"""

import logging

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from pathlib import Path

from src.renderers.base import BaseRenderer, RenderContext
from src.config.schemas import CodeBlockData
from src.utils.formatting import parse_inline_formatting
from src.utils.docx_utils import optimize_invisible_table, get_alignment_enum, fix_table_position

logger = logging.getLogger(__name__)


class CodeBlockRenderer(BaseRenderer):
    """
    Renderer for code block content nodes.
    
    Handles:
    - Monospace font code display
    - Optional caption above code
    - Code block specific formatting (no indent, single spacing)
    - If caption is present, uses a borderless table to repeat caption on new pages
    """
    
    @property
    def node_type(self) -> str:
        """Return the node type this renderer handles."""
        return "code"
    
    def render(self, context: RenderContext, data: CodeBlockData) -> None:
        """
        Render code block content to document.
        
        Args:
            context: Render context with container and config.
            data: Validated CodeBlockData model.
        """
        resolved_code = self._resolve_code_content(context, data)
        
        if data.caption:
            self._render_with_table(context, data, resolved_code)
        else:
            self._fill_code_container(context.container, resolved_code, context, is_cell=False)
            
    def _resolve_code_content(self, context: RenderContext, data: CodeBlockData) -> str:
        """
        Resolve code content from given data, handling file paths.
        
        Args:
            context: Current render context.
            data: Code block data containing code or path.
            
        Returns:
            Resolved code string.
        """
        if data.path:
            if data.code:
                logger.warning("CodeBlock: Both 'code' and 'path' are provided. 'path' takes precedence.")
                
            path_obj = Path(data.path)
            if not path_obj.is_absolute():
                file_path = context.resource_path / path_obj
            else:
                file_path = path_obj
                
            try:
                return file_path.read_text(encoding="utf-8")
            except Exception as e:
                logger.warning(f"Failed to read code file from {file_path}: {e}")
                return f"// WARNING: Failed to load code from {file_path}"
                
        return data.code or ""
            
    def _render_with_table(self, context: RenderContext, data: CodeBlockData, resolved_code: str) -> None:
        """Render caption and code inside a borderless 1x2 table."""
        table = context.container.add_table(rows=2, cols=1)
        fix_table_position(table)
        table.autofit = False
        
        # Determine width
        item_width = Cm(16.5)
        try:
            section = context.doc.sections[0]
            available = section.page_width - section.left_margin - section.right_margin
            if available > 0:
                item_width = available
        except Exception:
            pass
            
        table.columns[0].width = item_width
        optimize_invisible_table(table)
        
        # Row 0: Caption
        self._mark_header_row(table.rows[0])
        cell_caption = table.cell(0, 0)
        self._fill_caption_paragraph(cell_caption.paragraphs[0], data, context)
        
        # Row 1: Code
        cell_code = table.cell(1, 0)
        self._fill_code_container(cell_code, resolved_code, context, is_cell=True)

    def _fill_caption_paragraph(self, p, data: CodeBlockData, context: RenderContext) -> None:
        """
        Fill paragraph with caption formatting from global config.
        
        Uses 'caption' style for consistent look across all document objects
        (tables, listings, etc.).
        """
        style_config = context.config.styles.caption
        fonts = context.config.fonts
        
        # Apply paragraph-level formatting
        pf = p.paragraph_format
        pf.alignment = get_alignment_enum(style_config.alignment)
        pf.line_spacing = style_config.line_spacing
        pf.first_line_indent = Cm(style_config.first_line_indent_cm)
        pf.space_before = Pt(style_config.space_before_pt)
        pf.space_after = Pt(style_config.space_after_pt)
        
        # Reset indents to 0 unless specified otherwise
        pf.left_indent = Pt(0)
        pf.right_indent = Pt(0)
        
        parse_inline_formatting(
            paragraph=p,
            text=data.caption,
            default_font=fonts.default_name,
            custom_font=style_config.font_name,
            code_font=context.config.styles.inline_code.font_name if context.config.styles.inline_code else fonts.code_name,
            base_size_pt=style_config.font_size_pt
        )

    def _fill_code_container(self, container, code_text: str, context: RenderContext, is_cell: bool = False) -> None:
        """Fill container (cell or document) with code content as multiple paragraphs."""
        fonts = context.config.fonts
        code_style = context.config.styles.code_block
        
        lines = code_text.split('\n')
        # Prevent last empty newline from creating an extra paragraph gap
        if lines and lines[-1] == '':
            lines = lines[:-1]
            
        for i, line in enumerate(lines):
            if is_cell and i == 0:
                p = container.paragraphs[0]
            else:
                p = container.add_paragraph()
                
            pf = p.paragraph_format
            pf.line_spacing = code_style.line_spacing
            
            # Only add space_before to the VERY FIRST paragraph, and space_after to the VERY LAST
            pf.space_before = Pt(code_style.space_before_pt) if i == 0 else Pt(0)
            pf.space_after = Pt(code_style.space_after_pt) if i == len(lines)-1 else Pt(0)
            
            # Indents
            pf.first_line_indent = Pt(0)
            pf.left_indent = Pt(0)
            pf.right_indent = Pt(0)
            
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Don't add empty run for empty line to save XML space
            if line:
                run = p.add_run(line)
                run.font.name = code_style.font_name or fonts.code_name
                run.font.size = Pt(code_style.font_size_pt)

    def _mark_header_row(self, row) -> None:
        """Mark row to repeat on page breaks using OXML w:tblHeader."""
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)
