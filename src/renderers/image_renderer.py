"""
Image Renderer for Reports-Formater.

Renders images with optional captions and alignment.
"""

import io
import hashlib
import logging
from pathlib import Path

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from docx.shared import Cm, Mm, Emu, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL

from src.renderers.base import BaseRenderer, RenderContext
from src.config.schemas import ImageData
from src.utils.formatting import parse_inline_formatting
from src.utils.docx_utils import optimize_invisible_table

logger = logging.getLogger(__name__)

ALIGNMENT_MAP = {
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
}


class ImageRenderer(BaseRenderer):
    """
    Renderer for image content nodes.
    
    Handles:
    - Image insertion with width/height control
    - Image alignment
    - Optional captions below image
    """
    
    @property
    def node_type(self) -> str:
        """Return the node type this renderer handles."""
        return "image"
    
    def render(self, context: RenderContext, data: ImageData) -> None:
        """
        Render image content to document.
        """
        # Determine Max Width dynamically
        # Calculate available width based on page setup
        page_setup = context.config.page_setup
        # Word default A4 is 21.0 cm width
        # available_width = total_width - (left_margin + right_margin)
        total_page_width_cm = 21.0 
        available_width_cm = total_page_width_cm - (page_setup.margin_left_cm + page_setup.margin_right_cm)
        
        item_width = Cm(available_width_cm)
        
        try:
            # Try to get actual width from first section if available
            section = context.doc.sections[0]
            actual_avail_width = section.page_width - section.left_margin - section.right_margin
            if actual_avail_width > 0:
                item_width = actual_avail_width
        except Exception:
            logger.debug(f"ImageRenderer: Using calculated fallback width {available_width_cm}cm")
            
        # Determine image source
        image_to_insert = None
        missing = False
        
        if getattr(data, 'placeholder', False):
            image_to_insert = self._generate_placeholder_image("IMAGE PLACEHOLDER", '#fff2cc', '#d6b656')
        else:
            image_path = self._resolve_path(data.path, context.resource_path)
            if not image_path.exists():
                logger.error(f"ImageRenderer: Image not found: {image_path}")
                image_to_insert = self._generate_placeholder_image(f"MISSING IMAGE:\n{data.path}", '#ffebe6', 'red')
                missing = True
            else:
                image_to_insert = str(image_path)
                
        # Persist generated placeholder to temp file if necessary
        if isinstance(image_to_insert, io.BytesIO):
            temp_dir = context.resource_path / ".temp_images"
            temp_dir.mkdir(exist_ok=True)
            name_hash = hashlib.md5(str(data.path).encode()).hexdigest()
            temp_file = temp_dir / f"placeholder_{name_hash}.png"
            temp_file.write_bytes(image_to_insert.getvalue())
            image_to_insert = str(temp_file.resolve())

        # Apply fit_to_page logic dynamically constraining target item width
        if getattr(data, 'fit_to_page', False):
            try:
                from PIL import Image as PILImage
                with PILImage.open(image_to_insert) as img:
                    w_px, h_px = img.size
                if h_px > 0:
                    aspect_ratio = w_px / h_px
                    
                    # Determine total page height (A4 fallback is 29.7 cm)
                    page_height_cm = 29.7
                    try:
                        section = context.doc.sections[0]
                        if hasattr(section, 'page_height') and getattr(section, 'page_height') > 0:
                            if hasattr(section.page_height, 'cm'):
                                page_height_cm = section.page_height.cm
                            else:
                                page_height_cm = section.page_height / 360000.0 # 360000 EMUs per cm
                    except Exception:
                        pass
                        
                    max_height_cm = page_height_cm - page_setup.margin_top_cm - page_setup.margin_bottom_cm - page_setup.image_fit_padding_cm
                    
                    current_width_cm = item_width.cm if hasattr(item_width, 'cm') else item_width / 360000.0
                    projected_height_cm = current_width_cm / aspect_ratio
                    
                    if projected_height_cm > max_height_cm and max_height_cm > 0:
                        adjusted_width_cm = max_height_cm * aspect_ratio
                        item_width = Cm(adjusted_width_cm)
                        logger.debug(f"ImageRenderer: fit_to_page active. Adjusted width to {adjusted_width_cm:.2f}cm representing {max_height_cm:.2f}cm height.")
            except Exception as e:
                logger.warning(f"ImageRenderer: Failed to calculate fit_to_page constraints: {e}")

        # Create Layout Table (Single row for both to avoid LibreOffice keep_with_next bugs)
        table = context.container.add_table(rows=1, cols=1)
        table.autofit = False
        table.columns[0].width = item_width
        optimize_invisible_table(table)
        
        # Populate Image Cell (Row 0)
        cell = table.cell(0, 0)
        p_img = cell.paragraphs[0]
        # remove spacing for clean table fit
        p_img.paragraph_format.left_indent = Pt(0)
        p_img.paragraph_format.right_indent = Pt(0)
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(0)
        p_img.paragraph_format.first_line_indent = Pt(0)
        
        p_img.alignment = ALIGNMENT_MAP.get(data.align, WD_ALIGN_PARAGRAPH.CENTER)
        p_img.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        run_img = p_img.add_run()
        
        try:
            if data.width_cm:
                shape = run_img.add_picture(image_to_insert, width=Cm(data.width_cm))
            elif data.height_cm:
                shape = run_img.add_picture(image_to_insert, height=Cm(data.height_cm))
            else:
                shape = run_img.add_picture(image_to_insert, width=item_width)
                
            # IMPORTANT: LibreOffice Custom Image Margin Fix
            inline = shape._inline
            inline.set('distT', "0")
            inline.set('distB', "0")
            inline.set('distL', "0")
            inline.set('distR', "0")
        except Exception as e:
            logger.error(f"ImageRenderer: Failed to insert image: {e}")
            return
            
        # Add caption if provided (Same Cell, next paragraph)
        if data.caption:
            p_caption = cell.add_paragraph()
            p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_caption.paragraph_format.left_indent = Pt(0)
            p_caption.paragraph_format.right_indent = Pt(0)
            p_caption.paragraph_format.space_before = Pt(0)
            p_caption.paragraph_format.space_after = Pt(0)
            p_caption.paragraph_format.first_line_indent = Pt(0)
            
            parse_inline_formatting(
                paragraph=p_caption,
                text=data.caption,
                default_font=context.config.fonts.default_name,
                code_font=context.config.fonts.code_name,
                first_line_indent=0.0
            )

    @staticmethod
    def _generate_placeholder_image(text: str, bg_color: str, edge_color: str) -> io.BytesIO:
        """Generate a placeholder image for missing or explicitly placeholder images."""
        fig = plt.figure(figsize=(6, 2), dpi=100)
        fig.text(
            0.5, 0.5, text, 
            fontsize=12, color=edge_color, family='monospace',
            ha='center', va='center',
            bbox=dict(facecolor=bg_color, edgecolor=edge_color, boxstyle='square,pad=1.0', linewidth=2)
        )
        plt.axis('off')
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight', transparent=True)
        plt.close(fig)
        buf.seek(0)
        return buf

    def _resolve_path(self, path_str: str, resource_path: Path) -> Path:
        """Resolve image path."""
        path = Path(path_str)
        if path.is_absolute():
            return path
        return resource_path / path
