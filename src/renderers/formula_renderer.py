"""
Formula Renderer for Reports-Formater.

Renders LaTeX formulas as images using a Hybrid Strategy (Matplotlib/System LaTeX),
then inserts them into an invisible layout table (1 row, 2 columns) to ensure
strict positioning of the formula (Center) and its caption (Right) according to DSTU.
"""

import hashlib
import io
import logging
import shutil
import subprocess
import traceback
from typing import Tuple

import matplotlib
# Use 'Agg' backend to prevent GUI windows from popping up during rendering
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from PIL import Image

from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from src.renderers.base import BaseRenderer, RenderContext
from src.config.schemas import FormulaData
from src.utils.docx_utils import optimize_invisible_table

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Configuration Constants
# ---------------------------------------------------------------------------
RENDER_DPI = 300
"""DPI used by Matplotlib to rasterise the formula. Higher = crisper."""

INCHES_TO_CM = 2.54
"""Standard conversion factor: 1 inch = 2.54 cm."""

# ---------------------------------------------------------------------------
# Layout Constants (Invisible Table Strategy)
# ---------------------------------------------------------------------------
# Total width target: ~16.5 cm (A4 width 21cm - 2.5cm left - 1.5cm right margins)
# We reserve fixed space for the label to ensure it doesn't wrap awkwardly.
CAPTION_COLUMN_WIDTH_CM = 2.5
"""Width reserved for the right column containing the formula label (e.g., '(1.1)')."""

FORMULA_COLUMN_WIDTH_CM = 14.0
"""Width for the formula column. Images wider than this will be scaled down."""

# ---------------------------------------------------------------------------
# LaTeX Rendering Defaults
# ---------------------------------------------------------------------------
LATEX_FONT_SIZE = 14
"""Font size (pt) passed to Matplotlib's fig.text() for LaTeX rendering."""

SAVEFIG_PAD_INCHES = 0.02
"""Padding around the tight bounding box when saving the figure."""


class FormulaRenderer(BaseRenderer):
    """Renderer for formula content nodes.

    Architecture:
    1.  **Render:** Converts LaTeX to PNG using a Hybrid Strategy (Internal -> System).
    2.  **Layout:** Places the PNG and Caption into a borderless 1x2 Table.
        -   Cell 0: Formula (Centered, Scaled to fit).
        -   Cell 1: Caption (Right-aligned, Vertically centered).
    """

    @property
    def node_type(self) -> str:
        return "formula"

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def render(self, context: RenderContext, data: FormulaData) -> None:
        """Render a single formula node into the document.

        Args:
            context: Current rendering context (container, resource path, …).
            data:    Validated formula data (content, optional caption, align).

        Raises:
            OSError:  If the temporary image file cannot be written.
            RuntimeError: If image generation fails completely.
        """
        try:
            self._insert_formula_as_table(context, data)
        except (OSError, RuntimeError, ValueError) as exc:
            logger.error(
                "FormulaRenderer: Failed to render formula '%s': %s",
                data.content, exc,
            )
            traceback.print_exc()
            # Fallback: Insert error text directly into the document
            fallback = context.container.add_paragraph()
            run_err = fallback.add_run(f"[FORMULA ERROR: {data.content}]")
            run_err.font.name = context.config.fonts.default_name
            run_err.font.color.rgb = (255, 0, 0)

    # ------------------------------------------------------------------
    # Layout Logic (Invisible Table)
    # ------------------------------------------------------------------

    def _insert_formula_as_table(self, context: RenderContext, data: FormulaData) -> None:
        """Insert formula and caption using a 1x2 borderless table.

        Structure:
            | [Formula Image (Center)] | [Caption (Right)] |

        Args:
            context: Current rendering context.
            data:    Validated formula data.
        """
        # Generate Image (Hybrid Strategy)
        buf = self._render_latex_to_image(data.content)
        image_bytes = buf.getvalue()

        # Calculate Dimensions with Safety Scaling
        width_cm, height_cm = self._calculate_fitted_dimensions(buf)

        # Persist image to temp file
        temp_dir = context.resource_path / ".temp_formulas"
        temp_dir.mkdir(exist_ok=True)
        
        name_hash = hashlib.md5(data.content.encode()).hexdigest()
        temp_file = temp_dir / f"formula_{name_hash}.png"
        temp_file.write_bytes(image_bytes)

        # Create Layout Table
        table = context.container.add_table(rows=1, cols=2)
        
        # CRITICAL FIX: To force exact column widths in Word, we must:
        # a) Disable autofit.
        # b) Set the width on the COLUMN objects (not just cells).
        # This writes the <w:tblGrid> XML element which strictly defines layout.
        table.autofit = False 
        table.columns[0].width = Cm(FORMULA_COLUMN_WIDTH_CM)
        table.columns[1].width = Cm(CAPTION_COLUMN_WIDTH_CM)

        # Optimize invisible table
        optimize_invisible_table(table)

        # 5. Fill Formula Cell (Left)
        cell_formula = table.cell(0, 0)
        cell_formula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Ensure cell width is explicitly set (redundancy for safety)
        cell_formula.width = Cm(FORMULA_COLUMN_WIDTH_CM)

        p_formula = cell_formula.paragraphs[0]
        p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Reset indents to prevent global styles from shifting the image
        p_formula.paragraph_format.left_indent = Pt(0)
        p_formula.paragraph_format.right_indent = Pt(0)
        p_formula.paragraph_format.space_before = Pt(0)
        p_formula.paragraph_format.space_after = Pt(0)
        p_formula.paragraph_format.first_line_indent = Pt(0)
        
        run_formula = p_formula.add_run()
        run_formula.add_picture(
            str(temp_file.resolve()),
            width=Cm(width_cm),
            height=Cm(height_cm),
        )

        # Fill Caption Cell (Right)
        if data.caption:
            cell_caption = table.cell(0, 1)
            cell_caption.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Ensure cell width redundancy
            cell_caption.width = Cm(CAPTION_COLUMN_WIDTH_CM)
            
            p_caption = cell_caption.paragraphs[0]
            p_caption.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Hard coded. Maybe need to move it into config. 
            p_caption.paragraph_format.left_indent = Pt(0)
            p_caption.paragraph_format.right_indent = Pt(0)
            p_caption.paragraph_format.space_before = Pt(0)
            p_caption.paragraph_format.space_after = Pt(0)
            p_caption.paragraph_format.first_line_indent = Pt(0)
            
            # Apply header_footer style for consistency (since formula captions are like page numbers/labels)
            # or maybe normal caption style? Let's use caption style as per table renderer.
            cap_style = context.config.styles.caption
            fonts = context.config.fonts
            
            run_cap = p_caption.add_run(data.caption)
            run_cap.font.name = cap_style.font_name or fonts.default_name
            run_cap.font.size = Pt(cap_style.font_size_pt)

    def _calculate_fitted_dimensions(self, image_buffer: io.BytesIO) -> Tuple[float, float]:
        """Calculates image dimensions (cm), scaling down if it exceeds column width.

        Args:
            image_buffer: Buffer containing the image data.

        Returns:
            Tuple (width_cm, height_cm) ready for insertion.
        """
        image_buffer.seek(0)
        with Image.open(image_buffer) as pil_img:
            width_px, height_px = pil_img.size
        image_buffer.seek(0)

        # Convert raw pixels to cm based on render DPI
        width_cm = (width_px / RENDER_DPI) * INCHES_TO_CM
        height_cm = (height_px / RENDER_DPI) * INCHES_TO_CM

        # Check for overflow
        # We subtract a small buffer (0.1cm) to account for internal cell padding
        max_width = FORMULA_COLUMN_WIDTH_CM - 0.1
        
        if width_cm > max_width:
            scale_factor = max_width / width_cm
            width_cm = max_width
            height_cm = height_cm * scale_factor
            logger.info(
                "FormulaRenderer: Image too wide (%.2f cm). Scaled down to %.2f cm.", 
                (width_px / RENDER_DPI) * INCHES_TO_CM, width_cm
            )

        return width_cm, height_cm

    # ------------------------------------------------------------------
    # Image Generation (Hybrid Strategy)
    # ------------------------------------------------------------------

    @staticmethod
    def _render_latex_to_image(latex: str) -> io.BytesIO:
        """Render a LaTeX string to an in‑memory PNG image.

        Hybrid Strategy:
        1. Try Matplotlib's internal engine (fast, no deps).
        2. If syntax is unsupported (e.g. 'cases'), try system 'latex' (if available).
        3. If both fail, return an error placeholder image.

        Args:
            latex: Raw LaTeX expression (without ``$`` delimiters).

        Returns:
            A ``BytesIO`` buffer positioned at the start, containing the PNG.
        """
        # Attempt 1: Internal Matplotlib Engine (MathText)
        try:
            return FormulaRenderer._render_internal(latex)
        except ValueError:
            logger.info("Internal renderer failed on '%s'. Trying system LaTeX...", latex[:20])
        except Exception as e:
            logger.warning("Unexpected error in internal renderer: %s", e)

        # Attempt 2: System LaTeX (External)
        if FormulaRenderer._is_system_latex_available():
            try:
                return FormulaRenderer._render_external(latex)
            except Exception as e:
                logger.error("System LaTeX failed: %s", e)
        
        # Attempt 3: Graceful Failure (Error Placeholder)
        return FormulaRenderer._generate_error_image(
            "Formula Error:\nSyntax unsupported\nor LaTeX missing."
        )

    @staticmethod
    def _render_internal(latex: str) -> io.BytesIO:
        """Render using Matplotlib's built-in mathtext engine."""
        fig = plt.figure(figsize=(0.1, 0.1), dpi=RENDER_DPI)
        
        # Force internal engine
        with plt.rc_context({'text.usetex': False}):
            fig.text(0, 0, f"${latex}$", fontsize=LATEX_FONT_SIZE)
            
            buf = io.BytesIO()
            plt.axis('off')
            plt.savefig(
                buf,
                format='png',
                bbox_inches='tight',
                pad_inches=SAVEFIG_PAD_INCHES,
                dpi=RENDER_DPI,
                transparent=True,
            )
        plt.close(fig)
        buf.seek(0)
        return buf

    @staticmethod
    def _render_external(latex: str) -> io.BytesIO:
        """Render using system 'latex' binary via Matplotlib."""
        fig = plt.figure(figsize=(0.1, 0.1), dpi=RENDER_DPI)
        
        # Enable system latex and standard packages
        rc_params = {
            'text.usetex': True,
            'text.latex.preamble': r'\usepackage{amsmath} \usepackage{amssymb}'
        }
        
        with plt.rc_context(rc_params):
            fig.text(0, 0, f"${latex}$", fontsize=LATEX_FONT_SIZE)
            
            buf = io.BytesIO()
            plt.axis('off')
            plt.savefig(
                buf,
                format='png',
                bbox_inches='tight',
                pad_inches=SAVEFIG_PAD_INCHES,
                dpi=RENDER_DPI,
                transparent=True,
            )
        plt.close(fig)
        buf.seek(0)
        return buf

    @staticmethod
    def _is_system_latex_available() -> bool:
        """Check if 'latex' and a ghostscript/dvipng backend are present."""
        return (shutil.which('latex') is not None) and \
               ((shutil.which('dvipng') is not None) or (shutil.which('gs') is not None))

    @staticmethod
    def _generate_error_image(text: str) -> io.BytesIO:
        """Generate a placeholder image for failed formulas."""
        fig = plt.figure(figsize=(4, 1), dpi=100)
        fig.text(
            0.5, 0.5, text, 
            fontsize=10, color='red', family='monospace',
            ha='center', va='center',
            bbox=dict(facecolor='#ffebe6', edgecolor='red', boxstyle='round,pad=0.5')
        )
        plt.axis('off')
        buf = io.BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf