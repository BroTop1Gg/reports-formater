"""
Appendix Marker Renderer for Reports-Formater.

Renders appendix section markers (Додаток А, Додаток Б, etc.)
Creates new Word sections with independent headers/footers.
"""

import logging

from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.renderers.base import BaseRenderer, RenderContext
from src.config.schemas import AppendixMarkerData
from src.utils.docx_utils import add_page_number_field

logger = logging.getLogger(__name__)


class AppendixMarkerRenderer(BaseRenderer):
    """
    Renderer for appendix marker content nodes.
    
    Handles:
    - Creating new Word sections (WD_SECTION.NEW_PAGE)
    - Breaking header/footer links with previous section
    - Setting up independent page numbering for appendices
    - Rendering appendix title with Heading 1 style
    """
    
    @property
    def node_type(self) -> str:
        """Return the node type this renderer handles."""
        return "appendix_marker"
    
    def render(self, context: RenderContext, data: AppendixMarkerData) -> None:
        """
        Render appendix marker to document.
        
        Creates a new section with independent headers/footers,
        then renders the appendix title.
        
        Args:
            context: Render context with container and config.
            data: Validated AppendixMarkerData model.
        """
        # Create new section
        new_section = context.doc.add_section(WD_SECTION.NEW_PAGE)
        
        # Enable different first page header/footer
        new_section.different_first_page_header_footer = True
        
        # Break links with previous section
        new_section.header.is_linked_to_previous = False
        new_section.first_page_header.is_linked_to_previous = False
        
        # Configure first page header (page number top-right)
        self._setup_first_page_header(new_section, context)
        
        # Configure subsequent pages header (page number + continuation text)
        self._setup_continuation_header(new_section, context, data.label)
        
        # Render appendix title
        self._render_appendix_title(context, data)
        
        logger.info(f"AppendixMarkerRenderer: Created appendix section for '{data.label}'")
    
    def _setup_first_page_header(self, section, context: RenderContext) -> None:
        """
        Set up first page header with page number (top-right).
        
        If page_numbering.enabled is False, header remains empty.
        
        Args:
            section: The new Word section.
            context: Render context with config.
        """
        header = section.first_page_header
        
        # Clear existing content
        for p in header.paragraphs:
            p.clear()
        
        # If page numbering is disabled, leave header empty
        if not context.config.page_numbering.enabled:
            return
        
        # Get or create paragraph
        if len(header.paragraphs) == 0:
            p = header.add_paragraph()
        else:
            p = header.paragraphs[0]
        
        # Right alignment
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add page number field
        run = add_page_number_field(p)
        
        # Apply font formatting from config
        hf_style = context.config.styles.header_footer
        fonts = context.config.fonts
        run.font.name = hf_style.font_name or fonts.default_name
        run.font.size = Pt(hf_style.font_size_pt)
        run.bold = hf_style.bold
    
    def _setup_continuation_header(
        self, section, context: RenderContext, label: str
    ) -> None:
        """
        Set up continuation header with page number and "Продовження додатка X".
        
        If page_numbering.enabled is False, only "Продовження додатка X" is shown
        (centered, single paragraph).
        
        Args:
            section: The new Word section.
            context: Render context with config.
            label: Appendix letter label (А, Б, В, etc.).
        """
        header = section.header
        
        # Clear existing content
        for p in header.paragraphs:
            p.clear()
        
        hf_style = context.config.styles.header_footer
        fonts = context.config.fonts
        
        if context.config.page_numbering.enabled:
            # Paragraph 1: Page number (right-aligned)
            if len(header.paragraphs) == 0:
                p1 = header.add_paragraph()
            else:
                p1 = header.paragraphs[0]
            
            p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_num = add_page_number_field(p1)
            run_num.font.name = hf_style.font_name or fonts.default_name
            run_num.font.size = Pt(hf_style.font_size_pt)
            run_num.bold = hf_style.bold
            
            # Paragraph 2: Continuation text (centered)
            p2 = header.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_text = p2.add_run(f"Продовження додатка {label}")
            run_text.font.name = hf_style.font_name or fonts.default_name
            run_text.font.size = Pt(hf_style.font_size_pt)
            run_text.bold = hf_style.bold
        else:
            # Single centered paragraph with continuation text only
            if len(header.paragraphs) == 0:
                p1 = header.add_paragraph()
            else:
                p1 = header.paragraphs[0]
            
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_text = p1.add_run(f"Продовження додатка {label}")
            run_text.font.name = hf_style.font_name or fonts.default_name
            run_text.font.size = Pt(hf_style.font_size_pt)
            run_text.bold = hf_style.bold
    
    def _render_appendix_title(
        self, context: RenderContext, data: AppendixMarkerData
    ) -> None:
        """
        Render appendix title paragraph with Heading 1 style.
        
        Args:
            context: Render context with container.
            data: Appendix marker data with label and optional title.
        """
        # Build title text
        title_line = f"ДОДАТОК {data.label}"
        if data.title:
            title_text = f"{title_line}\n{data.title}"
        else:
            title_text = title_line
        
        # Add paragraph with Heading 1 style for TOC
        p = context.container.add_paragraph(title_text)
        
        # Try to apply Heading 1 style
        style_name = context.style_manager.get_style_name("Heading 1", fallback="Normal")
        try:
            p.style = style_name
        except (KeyError, ValueError) as e:
            logger.debug(f"Could not apply style '{style_name}': {e}")
        
        # Apply heading_1 config formatting
        style_config = context.config.styles.heading_1 or context.config.styles.heading_base
        fonts = context.config.fonts
        
        for run in p.runs:
            run.font.name = style_config.font_name or fonts.default_name
            run.font.size = Pt(style_config.font_size_pt)
            run.bold = style_config.bold
        
        # Apply paragraph formatting
        from docx.shared import Cm
        from src.utils.docx_utils import get_alignment_enum
        
        pf = p.paragraph_format
        pf.space_before = Pt(style_config.space_before_pt)
        pf.space_after = Pt(style_config.space_after_pt)
        pf.alignment = get_alignment_enum(style_config.alignment)
        pf.first_line_indent = Cm(style_config.first_line_indent_cm)
