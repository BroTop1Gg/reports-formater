"""
Report Factory for Reports-Formater.

Main entry point for building reports. Wires together all services.
"""

import logging
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

from src.config.loader import ConfigLoader
from src.config.models import ReportConfig
from src.services.style_manager import StyleManager
from src.services.placeholder_service import PlaceholderService
from src.services.rendering_service import RenderingService
from src.utils.file_io import FailSafeSaver
from src.utils.docx_utils import get_alignment_enum


# Import all renderers
from src.renderers.paragraph_renderer import ParagraphRenderer
from src.renderers.heading_renderer import HeadingRenderer
from src.renderers.list_renderer import ListRenderer
from src.renderers.table_renderer import TableRenderer
from src.renderers.image_renderer import ImageRenderer
from src.renderers.code_block_renderer import CodeBlockRenderer
from src.renderers.break_renderer import BreakRenderer
from src.renderers.formula_renderer import FormulaRenderer

logger = logging.getLogger(__name__)


class ReportFactory:
    """
    Factory for building Word reports from YAML content.
    
    Orchestrates the complete report generation process:
    1. Load configuration
    2. Load document template
    3. Replace placeholders
    4. Render content nodes
    5. Save output
    """
    
    def __init__(
        self,
        config_path: Optional[Path] = None,
        template_path: Optional[Path] = None,
    ):
        """
        Initialize ReportFactory.
        
        Args:
            config_path: Path to report_styles.json.
            template_path: Path to template DOCX file.
        """
        self._config_loader = ConfigLoader(config_path)
        self._template_path = template_path
        self._saver = FailSafeSaver()
        
        # Initialize rendering service with all renderers
        self._rendering_service = RenderingService()
        self._rendering_service.register_all([
            ParagraphRenderer(),
            HeadingRenderer(),
            ListRenderer(),
            TableRenderer(),
            ImageRenderer(),
            CodeBlockRenderer(),
            BreakRenderer(),
            FormulaRenderer(),
        ])
    
    def build(
        self,
        yaml_data: Dict[str, Any],
        output_path: Path,
        resource_path: Optional[Path] = None,
    ) -> Path:
        """
        Build report from YAML data.
        
        Args:
            yaml_data: Parsed YAML content.
            output_path: Output file path.
            resource_path: Base path for resolving relative resources.
            
        Returns:
            Actual path where document was saved.
        """
        output_path = Path(output_path)
        resource_path = resource_path or output_path.parent
        
        config = self._config_loader.load(yaml_data)
        
        # Load document template
        # Support metadata-driven template selection (Requirement: metadata.template_path)
        metadata_config = yaml_data.get("metadata", {})
        yml_template = metadata_config.get("template_path")
        
        # Priority: CLI argument > YAML metadata > Default template
        effective_template = self._template_path
        if not effective_template and yml_template:
            effective_template = Path(yml_template)
            
        doc = self._get_document(effective_template)
        
        self._setup_page_layout(doc, config, yaml_data)
        
        # Replace placeholders with flattening support
        mapping = metadata_config.get("mapping", {})
        # Support both flat metadata and nested 'mapping' dict for compatibility
        flat_metadata = {
            k: v for k, v in metadata_config.items() 
            if k not in ["mapping", "template_path"]
        }
        flat_metadata.update(mapping)
        
        if flat_metadata:
            placeholder_service = PlaceholderService()
            placeholder_service.replace_all(doc, flat_metadata)
        
        style_manager = StyleManager(doc)
        
        # Create render context
        context = self._rendering_service.create_context(
            doc=doc,
            container=doc,
            config=config,
            style_manager=style_manager,
            resource_path=resource_path,
        )
        
        # Update dispatch to use proper context
        def dispatch_with_context(data: Any) -> None:
            self._rendering_service.dispatch(context, data)
        
        context = context.__class__(
            doc=context.doc,
            container=context.container,
            config=context.config,
            style_manager=context.style_manager,
            resource_path=context.resource_path,
            dispatch=dispatch_with_context,
            list_level=context.list_level,
        )

        # Clear the default empty paragraph before rendering new content
        self._clear_initial_content(doc)
        
        # Parse content nodes
        from src.config.schemas import parse_content_node
        from src.services.spacing_engine import SpacingEngine
        
        content_raw = yaml_data.get("content", [])
        
        parsed_nodes = []
        for c in content_raw:
            try:
                parsed_nodes.append(parse_content_node(c))
            except Exception as e:
                logger.error(f"ReportFactory: Failed to parse node {c}: {e}")
        
        # Apply data-driven spacing engine using configuration
        # .model_dump() converts Pydantic V2 models to dict. Fallback to .dict() for V1.
        try:
            rules_dict = config.spacing_rules.model_dump()
        except AttributeError:
            rules_dict = config.spacing_rules.dict()
            
        spacing_engine = SpacingEngine(rules_dict)
        spaced_nodes = spacing_engine.process(parsed_nodes)
        
        # Render content
        self._rendering_service.render_content(context, spaced_nodes)
        
        # Spare function call. Better to call this function here, or trying to find why is margins broke.
        self._finalize_margins(doc, config)
        
        # Save with fail-safe
        return self._saver.save(doc, output_path)
    
    def _get_document(self, template_path: Optional[Path]) -> Document:
        """
        Load document template or create blank.
        
        Args:
            template_path: Path to template DOCX file.
            
        Returns:
            Document instance.
        """
        if template_path and template_path.exists():
            logger.info(f"Loading template: {template_path}")
            return Document(str(template_path))
        
        # Try default template
        default_path = Path(__file__).parent / "DEFAULT_TEMPLATE.docx"
        if default_path.exists():
            logger.info(f"Loading default template: {default_path}")
            return Document(str(default_path))
        
        logger.warning(f"No template found at '{template_path}'. Using blank document.")
        return Document()
    
    def _setup_page_layout(
        self, 
        doc: Document, 
        config: ReportConfig,
        yaml_data: Dict[str, Any],
    ) -> None:
        """
        Apply page margins and headers/footers from configuration.
        
        Args:
            doc: Document to configure.
            config: Report configuration.
            yaml_data: YAML data for additional settings.
        """
        page = config.page_setup
        
        for section in doc.sections:
            section.top_margin = Cm(page.margin_top_cm)
            section.bottom_margin = Cm(page.margin_bottom_cm)
            section.left_margin = Cm(page.margin_left_cm)
            section.right_margin = Cm(page.margin_right_cm)
            
            # Set Header/Footer distances from page_setup
            section.header_distance = Cm(page.header_distance_cm)
            section.footer_distance = Cm(page.footer_distance_cm)
        
        # Process global overrides
        # Configure header and footer
        self._configure_header_footer(doc, config, yaml_data)
    
    def _configure_header_footer(
        self, 
        doc: Document, 
        config: ReportConfig, 
        yaml_data: Dict[str, Any]
    ) -> None:
        """
        Configures page numbering and optional text in header or footer.
        
        Logic ported from builder.py to ensure correct positioning and clearing.
        """
        # Get settings with YAML overrides having priority
        pn_cfg = config.page_numbering
        
        # Determine if enabled: YAML override > Config default
        if "page_numbering" in yaml_data:
            raw_pn = yaml_data["page_numbering"]
            if isinstance(raw_pn, bool):
                show_numbering = raw_pn
            elif isinstance(raw_pn, dict):
                show_numbering = raw_pn.get("enabled", True)
            else:
                show_numbering = False
        else:
            show_numbering = pn_cfg.enabled

        header_text = yaml_data.get("header_text")
        
        # Determine if we need to write anything
        has_content = show_numbering or header_text
        
        page = config.page_setup

        for section in doc.sections:
            # Identification of targets

            # Identification of targets
            all_headers = [section.header, section.first_page_header, section.even_page_header]
            all_footers = [section.footer, section.first_page_footer, section.even_page_footer]
            
            if not has_content:
                # Block explanation: We dont have better way to clear header and footer, and becouse we just walking through all sections, we need to clear all headers and footers the document and clear them.
                # I dont understand where is this shit (headers and footers) added, but we need to remove it if this not needed.
                # CLEAR-ONLY STRATEGY
                for target in all_headers + all_footers:
                    if target: self._clear_container(target)
            else:
                if pn_cfg.position == 'footer':
                    active_target = section.footer
                    for h in all_headers:
                        if h: self._clear_container(h)
                else:
                    active_target = section.header
                    for f in all_footers:
                        if f: self._clear_container(f)
                
                # If header/footer not initialized, skip
                if not active_target: continue

                # We grand that we work only with one paragraph in header/footer
                if len(active_target.paragraphs) == 0:
                    active_target.add_paragraph()
                
                p = active_target.paragraphs[0]
                p.clear()
                
                # Cleaning other paragraphs
                for i in range(1, len(active_target.paragraphs)):
                     active_target.paragraphs[i].text = ""
                     active_target.paragraphs[i].clear()

                # SMART ALIGNMENT LOGIC
                # If we have both Text and Number, user expects "Spread" (Text Left, Number Right)" like:
                # | SOME_TEXT                                PAGE_NUMBER |
                # And becouse we use here tabulating to do that.
                # We follow the 'Dumb Builder' + 'Consistent Style' philosophy.
                hf_style = config.styles.header_footer
                p.alignment = get_alignment_enum(hf_style.alignment)
                
                # 1. Handle "Spread" Alignment (Text LEFT, Page Number RIGHT) using Tab Stops
                if header_text and show_numbering:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.tab_stops.clear_all()
                    
                    # Tab position = Total Page Width - Margins
                    page_w = section.page_width
                    left_m = section.left_margin
                    right_m = section.right_margin
                    
                    if page_w and left_m and right_m:
                        tab_pos = page_w - left_m - right_m
                        p.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
                    else:
                        # Fallback A4 = 17cm (21 - 2.5 - 1.5)
                        p.paragraph_format.tab_stops.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.RIGHT)
                
                # 2. Add Header/Footer Text
                if header_text:
                    run_text = p.add_run(f"{header_text}")
                    run_text.font.name = hf_style.font_name or config.fonts.default_name
                    run_text.font.size = Pt(hf_style.font_size_pt)
                    run_text.bold = hf_style.bold
                    
                    # If we need to spread, add the tab now
                    if show_numbering:
                        run_text.add_tab()
                        run_text.add_tab()
                
                # 3. Add Page Number
                if show_numbering:
                    run_num = self._add_page_number_field(p, config)
                    run_num.font.name = hf_style.font_name or config.fonts.default_name
                    run_num.font.size = Pt(hf_style.font_size_pt)
                    run_num.bold = hf_style.bold

    def _clear_container(self, container):
        """Helper to clear all content from a Header/Footer container."""
        for p in container.paragraphs:
            p.text = ""
            p.clear()
            
        for t in container.tables:
             for row in t.rows:
                 for cell in row.cells:
                     for p in cell.paragraphs:
                         p.text = ""
                         p.clear()

    def _add_page_number_field(self, paragraph, config: ReportConfig):
        """Adds a PAGE number field to a paragraph."""
        run = paragraph.add_run()
        
        def create_element(name):
            return OxmlElement(name)

        def create_attribute(element, name, value):
            element.set(qn(name), value)

        # Field Structure: begin -> instrText (PAGE) -> separate -> end
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')
        run._r.append(fldChar1)
        
        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)
        
        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'separate')
        run._r.append(fldChar2)
        
        fldChar3 = create_element('w:fldChar')
        create_attribute(fldChar3, 'w:fldCharType', 'end')
        run._r.append(fldChar3)
        
        return run

    def _finalize_margins(self, doc: Document, config: ReportConfig) -> None:
        """
        Ensure margins persist after all operations.
        
        Args:
            doc: Document to finalize.
            config: Report configuration.
        """
        page = config.page_setup
        
        for section in doc.sections:
            section.top_margin = Cm(page.margin_top_cm)
            section.bottom_margin = Cm(page.margin_bottom_cm)
            section.left_margin = Cm(page.margin_left_cm)
            section.right_margin = Cm(page.margin_right_cm)

    def _clear_initial_content(self, doc: Document) -> None:
        """
        Removes the initial empty paragraph created by default in new documents.
        This prevents an unwanted empty line at the very top of the report.
        """
        # A new document from template usually has exactly one empty paragraph.
        # We remove it only if it is indeed empty to avoid deleting template content (if any).
        if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text.strip():
            p = doc.paragraphs[0]
            p._element.getparent().remove(p._element)
            p._p = p._element = None