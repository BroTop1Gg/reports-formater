"""
Placeholder Service for Reports-Formater.

Replaces {{KEY}} placeholders in Word documents with values from metadata.
Implements Cascade Strategy for fragmented run handling.
"""

import logging
from enum import Enum
from typing import Dict, Any, Optional, List

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

logger = logging.getLogger(__name__)


class PlaceholderService:
    """
    Service for replacing placeholders in Word documents.
    
    Implements Cascade Strategy where it first attempts run-level replacement to
    preserve formatting, then falls back to paragraph-level replacement for
    fragmented runs. Finally, it handles cleanup of unused keys based on configuration.
    
    Handles MS Word's tendency to split {{KEY}} into multiple XML runs.
    """
    
    def __init__(self):
        """
        Initialize PlaceholderService.
        """
        self._replacement_count = 0
    
    def replace_all(
        self, 
        doc: Document, 
        metadata: Dict[str, Any],
    ) -> int:
        """
        Replace all placeholders in document.
        
        Searches body, headers, footers, and tables.
        
        Args:
            doc: Word Document object.
            metadata: Key-value pairs for replacement.
            
        Returns:
            Number of replacements made.
        """
        if not metadata:
            logger.debug("PlaceholderService: No metadata provided")
            return 0
        
        self._replacement_count = 0
        
        logger.debug(
            f"PlaceholderService: Starting replacement with {len(metadata)} keys"
        )
        
        # Replace in main body
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, metadata)
        
        for table in doc.tables:
            self._replace_in_table(table, metadata)
        
        # Replace in all headers and footers
        for section in doc.sections:
            parts = [
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ]
            
            for part in parts:
                try:
                    for paragraph in part.paragraphs:
                        self._replace_in_paragraph(paragraph, metadata)
                    for table in part.tables:
                        self._replace_in_table(table, metadata)
                except Exception:
                    # Some section parts may not be initialized
                    # Log this error? How error can be found if it just 'eated' by try-except with no actions?
                    pass
        
        logger.info(
            f"PlaceholderService: Completed with {self._replacement_count} replacements"
        )
        return self._replacement_count
    
    def _replace_in_paragraph(
        self, 
        paragraph: Paragraph, 
        metadata: Dict[str, Any],
    ) -> None:
        """
        Replace placeholders in a single paragraph.
        
        Implements Cascade Strategy. First, it attempts run-level replacement to
        preserve original formatting (bold, italic, etc.). If that fails due to
        fragmented runs, it falls back to paragraph-level replacement, which may
        lose mixed inline formatting but ensures the value is inserted.
        
        Args:
            paragraph: Paragraph object to process.
            metadata: Key-value pairs for replacement.
        """
        if not paragraph.text:
            return
        
        for key, value in metadata.items():
            # Support both original and uppercase keys
            target_placeholders = [f"{{{{{key}}}}}", f"{{{{{key.upper()}}}}}"]
            
            for placeholder in target_placeholders:
                if placeholder not in paragraph.text:
                    continue
                
                # Attempt run-level replacement (preserves formatting)
                run_replaced = False
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
                        run_replaced = True
                        self._replacement_count += 1
                        logger.debug(
                            f"PlaceholderService: Run-level replacement "
                            f"'{placeholder}' -> '{value}'"
                        )
                
                # Fallback to paragraph-level replacement for fragmented runs
                # IMPORTANT: The Writer often splits {{KEY}} into multiple runs
                # e.g., '{', '{', 'KEY', '}', '}' as separate runs
                if not run_replaced and placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        placeholder, 
                        str(value)
                    )
                    self._replacement_count += 1
                    logger.debug(
                        f"PlaceholderService: Paragraph-level fallback "
                        f"'{placeholder}' -> '{value}'"
                    )
    
    def _replace_in_table(
        self, 
        table: Table, 
        metadata: Dict[str, Any],
    ) -> None:
        """
        Replace placeholders in all table cells.
        
        Args:
            table: Table object to process.
            metadata: Key-value pairs for replacement.
        """
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_in_paragraph(paragraph, metadata)
    
    def cleanup_unused(
        self, 
        doc: Document, 
        used_keys: Optional[List[str]] = None,
    ) -> None:
        """
        Apply cleanup strategy to unused placeholders.
        
        Note: Currently, we follow a 'Dumb Builder' philosophy where unused 
        placeholders are left visible in the document to alert the user 
        of missing data. If automated cleanup strategies (like CLEAR or 
        HIGHLIGHT) are needed in the future, they should be implemented here.
        
        Args:
            doc: Word Document object.
            used_keys: Optional list of keys that were used.
        """
        pass
